import base64
import binascii
import html
import io
import json
import re
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

WORD_EXPORT_IMAGE_SCALE = 0.6
WORD_EXPORT_IMAGE_BASE_WIDTH_INCHES = 6.0
WORD_EXPORT_IMAGE_WIDTH_INCHES = WORD_EXPORT_IMAGE_BASE_WIDTH_INCHES * WORD_EXPORT_IMAGE_SCALE
WORD_EXPORT_FONT_LATIN = "Times New Roman"
WORD_EXPORT_FONT_EAST_ASIA = "Microsoft YaHei"
WORD_EXPORT_FONT_MONO = "Consolas"
WORD_EXPORT_TEXT_COLOR = RGBColor(31, 41, 55)
WORD_EXPORT_HEADING_COLOR = RGBColor(17, 24, 39)
WORD_EXPORT_MUTED_COLOR = RGBColor(107, 114, 128)
WORD_EXPORT_CODE_BACKGROUND = "F5F7FA"
WORD_EXPORT_QUOTE_BACKGROUND = "F8FAFC"
FIG_PLACEHOLDER_CORE_PATTERN = r"[\[\uFF3B\u3010]?\s*FIG\s*:?\s*\d+\s*[\]\uFF3D\u3011]?"
FIG_PLACEHOLDER_PATTERN = rf"(?<![A-Za-z0-9_]){FIG_PLACEHOLDER_CORE_PATTERN}(?![A-Za-z0-9_])"
FIG_TRAILING_PUNCTUATION_PATTERN = re.compile(
    rf"(?<![A-Za-z0-9_])(?P<placeholders>{FIG_PLACEHOLDER_CORE_PATTERN}(?:\s*{FIG_PLACEHOLDER_CORE_PATTERN})*)(?![A-Za-z0-9_])(?P<spacing>\s*)(?P<punctuation>[\u3002\uFF01\uFF1F\uFF1B\uFF0C\u3001\uFF1A]+)",
    flags=re.IGNORECASE,
)
INLINE_HEADING_BODY_STARTERS = (
    "本章",
    "本节",
    "本部分",
    "本报告",
    "本文",
    "首先",
    "其次",
    "然后",
    "接下来",
    "随后",
    "最后",
    "此外",
    "同时",
    "通过",
    "针对",
    "为了",
    "基于",
    "这里",
    "因此",
    "其中",
    "需要",
)
INLINE_HEADING_BODY_PATTERN = re.compile(
    rf"^(?P<title>.+?[。！？!?])\s*(?P<body>(?:{'|'.join(map(re.escape, INLINE_HEADING_BODY_STARTERS))}).+)$"
)
INLINE_MARKDOWN_HEADING_AFTER_TEXT_PATTERN = re.compile(
    r"^(?P<prefix>.*?[。！？!?；;：:])\s*(?P<heading>#{1,6}[ \t\u3000]*.+)$"
)


def maybe_json_loads(value: Any) -> Any:
    if not isinstance(value, str):
        return value

    stripped = value.strip()
    if not stripped:
        return value

    try:
        return json.loads(stripped)
    except json.JSONDecodeError:
        return value


def find_nested_field(data: Any, field_name: str) -> Any:
    if isinstance(data, dict):
        if field_name in data:
            return data[field_name]

        for nested_value in data.values():
            nested = find_nested_field(nested_value, field_name)
            if nested is not None:
                return nested

    if isinstance(data, list):
        for item in data:
            nested = find_nested_field(item, field_name)
            if nested is not None:
                return nested

    return None


def find_first_nested_field(data: Any, field_names: list[str]) -> Any:
    for field_name in field_names:
        value = find_nested_field(data, field_name)
        if value is not None:
            return value
    return None


def stringify_string(value: Any) -> str:
    value = maybe_json_loads(value)

    if value is None:
        return ""

    if isinstance(value, str):
        return value.strip()

    return json.dumps(value, ensure_ascii=False, indent=2)


def normalize_toc_list(value: Any) -> list[str]:
    parsed_value = maybe_json_loads(value)

    if isinstance(parsed_value, list):
        normalized_items: list[str] = []
        for item in parsed_value:
            item_text = str(item).replace("\\r\\n", "\n").replace("\\n", "\n").strip()
            if not item_text:
                continue
            normalized_items.extend([line.strip() for line in item_text.splitlines() if line.strip()])
        return normalized_items

    if isinstance(parsed_value, str):
        normalized_text = parsed_value.replace("\\r\\n", "\n").replace("\\n", "\n")
        return [line.strip() for line in normalized_text.splitlines() if line.strip()]

    return []


def extract_report_markdown(result: Any) -> str:
    if isinstance(result, str):
        return result.strip()

    candidate = find_first_nested_field(
        result,
        [
            "report_markdown",
            "markdown",
            "report_md",
            "md",
            "report_content",
            "content",
            "report",
            "body",
        ],
    )
    return stringify_string(candidate)


def extract_report_html(result: Any) -> str:
    candidate = find_first_nested_field(result, ["final_html", "report_html", "html"])
    if isinstance(candidate, list):
        return "".join(str(item) for item in candidate if item is not None).strip()
    return stringify_string(candidate)


def extract_report_text(result: Any) -> str:
    candidate = find_first_nested_field(
        result,
        ["report_text", "text", "report_content", "content", "report", "body"],
    )
    return stringify_string(candidate)


def extract_report_word_bytes(result: Any) -> bytes | None:
    candidate = find_first_nested_field(
        result,
        ["report_word", "word", "report_word_base64", "word_base64", "docx_base64"],
    )

    if isinstance(candidate, (bytes, bytearray)):
        return bytes(candidate)

    if not isinstance(candidate, str):
        return None

    stripped = candidate.strip()
    if not stripped:
        return None

    try:
        return base64.b64decode(stripped, validate=True)
    except (binascii.Error, ValueError):
        return None


def normalize_trailing_punctuation_before_figure_placeholder(text: str) -> str:
    if not isinstance(text, str) or not text.strip():
        return text

    return FIG_TRAILING_PUNCTUATION_PATTERN.sub(
        lambda match: f"{match.group('punctuation')}{match.group('spacing')}{match.group('placeholders')}",
        text,
    )


def _extract_markdown_heading_text(text: str) -> str | None:
    parsed_heading = _parse_markdown_heading_line(text)
    if parsed_heading is None:
        return None
    return parsed_heading[0]


def _split_inline_heading_content(text: str) -> tuple[str, str | None]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return "", None

    body_starters = (*INLINE_HEADING_BODY_STARTERS, "\u672c\u9636\u6bb5", "\u8be5\u9636\u6bb5")

    spaced_split = re.match(r"^(?P<title>\S(?:.*?\S)?)\s{2,}(?P<body>\S.*)$", text.strip())
    if spaced_split:
        return spaced_split.group("title").strip(), spaced_split.group("body").strip()

    starter_split = re.match(
        rf"^(?P<title>.+?)\s+(?P<body>(?:{'|'.join(map(re.escape, body_starters))}).+)$",
        normalized,
    )
    if starter_split:
        return starter_split.group("title").strip(), starter_split.group("body").strip()

    inline_split = INLINE_HEADING_BODY_PATTERN.match(normalized)
    if inline_split:
        return inline_split.group("title").strip(), inline_split.group("body").strip()

    return normalized, None


def _parse_markdown_heading_line(text: str) -> tuple[str, str | None] | None:
    if not isinstance(text, str):
        return None

    normalized = text.replace("\\r\\n", "\n").replace("\\n", "\n").strip()
    if not normalized:
        return None

    heading_match = re.match(r"^#{1,6}[ \t\u3000]*(.*)$", normalized)
    if not heading_match:
        return None

    heading_text = heading_match.group(1).strip()
    if not heading_text:
        return None

    title_text, body_text = _split_inline_heading_content(heading_text)
    if not title_text:
        return None
    return title_text, body_text


def _split_text_with_markdown_headings(text: str) -> list[tuple[str, str]]:
    normalized = text.strip()
    if not normalized:
        return []

    parsed_heading = _parse_markdown_heading_line(normalized)
    if parsed_heading:
        heading_text, body_text = parsed_heading
        segments: list[tuple[str, str]] = [("heading", heading_text)]
        if body_text:
            segments.append(("text", body_text))
        return segments

    inline_heading_match = INLINE_MARKDOWN_HEADING_AFTER_TEXT_PATTERN.match(normalized)
    if inline_heading_match:
        prefix_text = inline_heading_match.group("prefix").strip()
        heading_segments = _split_text_with_markdown_headings(inline_heading_match.group("heading"))
        if heading_segments:
            segments = []
            if prefix_text:
                segments.append(("text", prefix_text))
            segments.extend(heading_segments)
            return segments

    return [("text", normalized)]


def _split_markdown_heading_lines(text: str) -> list[tuple[str, str]]:
    if not isinstance(text, str):
        return []

    normalized = text.replace("\\r\\n", "\n").replace("\\n", "\n")
    lines = [line.strip() for line in normalized.split("\n") if line.strip()]
    if not lines:
        return []

    parsed_lines: list[tuple[str, str]] = []
    has_heading = False
    for line in lines:
        segments = _split_text_with_markdown_headings(line)
        parsed_lines.extend(segments)
        has_heading = has_heading or any(line_kind == "heading" for line_kind, _ in segments)

    return parsed_lines if has_heading else []


def html_to_markdown(html_text: str) -> str:
    if not html_text.strip():
        return ""

    html_text = normalize_trailing_punctuation_before_figure_placeholder(html_text)

    soup = BeautifulSoup(html_text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    root = soup.find("main") or soup.body or soup
    lines: list[str] = []

    block_tags = {
        "article",
        "aside",
        "div",
        "main",
        "section",
        "p",
        "ul",
        "ol",
        "li",
        "figure",
    }

    def add_line(text: str) -> None:
        clean = text.strip()
        if clean:
            lines.append(clean)

    def walk(node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                add_line(text)
            return

        if not isinstance(node, Tag):
            return

        if node.name in {"script", "style", "noscript"}:
            return

        if node.name and node.name.startswith("h") and len(node.name) == 2 and node.name[1].isdigit():
            text = node.get_text(" ", strip=True)
            if text:
                add_line(f"{'#' * int(node.name[1])} {text}")
            return

        if node.name == "li":
            text = node.get_text(" ", strip=True)
            if text:
                add_line(f"- {text}")
            return

        if node.name == "img":
            img_src = (node.get("src") or "").strip()
            alt_text = (node.get("alt") or "图表").strip()
            if img_src:
                add_line(f"![{alt_text}]({img_src})")
            return

        if node.name == "div" and "report-figure-block" in (node.get("class") or []):
            img = node.find("img")
            img_src = (img.get("src") or "").strip() if img else ""
            caption_tag = node.find(class_="report-figure-caption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else ""
            if img_src:
                add_line(f"![{caption_text or '图表'}]({img_src})")
            if caption_text:
                add_line(caption_text)
            return

        if node.name == "p":
            text = node.get_text(" ", strip=True)
            if text:
                add_line(text)
            return

        if node.name in block_tags:
            child_tags = [child for child in node.children if isinstance(child, Tag)]
            has_block_children = any(
                (
                    child.name in block_tags
                    or (child.name and child.name.startswith("h") and len(child.name) == 2 and child.name[1].isdigit())
                    or child.name == "img"
                )
                for child in child_tags
            )
            if not has_block_children:
                text = node.get_text(" ", strip=True)
                if text:
                    add_line(text)
                return

            direct_text = " ".join(
                text.strip()
                for text in node.find_all(string=True, recursive=False)
                if text and text.strip()
            )
            if direct_text:
                add_line(direct_text)
            for child in node.children:
                walk(child)
            return

        for child in node.children:
            walk(child)

    for child in root.children:
        walk(child)

    return "\n\n".join(lines).strip()


def build_markdown_preview_from_html(html_text: str, max_chars: int = 20000) -> str:
    if not html_text.strip():
        return ""

    soup = BeautifulSoup(html_text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    root = soup.find("main") or soup.body or soup
    blocks: list[str] = []

    def append_block(text: str) -> None:
        clean = re.sub(r"\s+", " ", text).strip()
        if clean:
            blocks.append(clean)

    def walk(node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                append_block(text)
            return

        if not isinstance(node, Tag):
            return

        if node.name in {"script", "style", "noscript"}:
            return

        if node.name == "br":
            blocks.append("")
            return

        if node.name and node.name.startswith("h") and len(node.name) == 2 and node.name[1].isdigit():
            text = node.get_text(" ", strip=True)
            if text:
                append_block(f"{'#' * int(node.name[1])} {text}")
            return

        if node.name == "li":
            text = node.get_text(" ", strip=True)
            if text:
                append_block(f"- {text}")
            return

        if node.name == "img":
            append_block("![图表](embedded-image)")
            return

        if node.name == "div" and "report-figure-block" in (node.get("class") or []):
            caption_tag = node.find(class_="report-figure-caption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else ""
            append_block(f"![{caption_text or '图表'}](embedded-image)")
            if caption_text:
                append_block(caption_text)
            return

        if node.name in {"p", "figcaption"}:
            text = node.get_text(" ", strip=True)
            if text:
                append_block(text)
            return

        for child in node.children:
            walk(child)

    for child in root.children:
        walk(child)

    preview = "\n\n".join(blocks)
    preview = re.sub(
        r"(?:\s*!\[图表\]\(embedded-image\)\s*){2,}",
        "\n\n![图表](embedded-image)\n\n",
        preview,
        flags=re.IGNORECASE,
    )
    preview = re.sub(r"\n{3,}", "\n\n", preview).strip()

    if len(preview) > max_chars:
        preview = preview[:max_chars].rstrip() + "\n\n...[预览已截断，下载文件保留完整内容]"

    return preview


def _configure_doc_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = WORD_EXPORT_FONT_LATIN
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = WORD_EXPORT_TEXT_COLOR
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(8)

    title_style = doc.styles["Title"]
    title_style.font.name = WORD_EXPORT_FONT_LATIN
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = WORD_EXPORT_HEADING_COLOR
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(18)
    title_style.paragraph_format.keep_with_next = True

    heading_sizes = {
        "Heading 1": 18,
        "Heading 2": 15,
        "Heading 3": 13,
        "Heading 4": 12,
        "Heading 5": 11,
        "Heading 6": 11,
    }
    for style_name, font_size in heading_sizes.items():
        heading_style = doc.styles[style_name]
        heading_style.font.name = WORD_EXPORT_FONT_LATIN
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
        heading_style.font.size = Pt(font_size)
        heading_style.font.bold = True
        heading_style.font.color.rgb = WORD_EXPORT_HEADING_COLOR
        heading_style.paragraph_format.space_before = Pt(16 if style_name == "Heading 1" else 14)
        heading_style.paragraph_format.space_after = Pt(6)
        heading_style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        list_style = doc.styles[style_name]
        list_style.font.name = WORD_EXPORT_FONT_LATIN
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
        list_style.font.size = Pt(11)
        list_style.font.color.rgb = WORD_EXPORT_TEXT_COLOR
        list_style.paragraph_format.line_spacing = 1.35
        list_style.paragraph_format.space_after = Pt(4)

    code_style = _get_or_create_paragraph_style(doc, "ChatGPT Code")
    code_style.base_style = normal_style
    code_style.font.name = WORD_EXPORT_FONT_MONO
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_MONO)
    code_style.font.size = Pt(9.5)
    code_style.font.color.rgb = WORD_EXPORT_TEXT_COLOR
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.right_indent = Inches(0.05)
    code_style.paragraph_format.line_spacing = 1.15
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    quote_style = _get_or_create_paragraph_style(doc, "ChatGPT Quote")
    quote_style.base_style = normal_style
    quote_style.font.name = WORD_EXPORT_FONT_LATIN
    quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    quote_style.font.size = Pt(10.5)
    quote_style.font.italic = True
    quote_style.font.color.rgb = WORD_EXPORT_MUTED_COLOR
    quote_style.paragraph_format.left_indent = Inches(0.25)
    quote_style.paragraph_format.right_indent = Inches(0.05)
    quote_style.paragraph_format.line_spacing = 1.35
    quote_style.paragraph_format.space_before = Pt(6)
    quote_style.paragraph_format.space_after = Pt(6)

    caption_style = _get_or_create_paragraph_style(doc, "ChatGPT Caption")
    caption_style.base_style = normal_style
    caption_style.font.name = WORD_EXPORT_FONT_LATIN
    caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_EXPORT_FONT_EAST_ASIA)
    caption_style.font.size = Pt(9.5)
    caption_style.font.italic = True
    caption_style.font.color.rgb = WORD_EXPORT_MUTED_COLOR
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.line_spacing = 1.2
    caption_style.paragraph_format.space_before = Pt(0)
    caption_style.paragraph_format.space_after = Pt(10)


def _get_or_create_paragraph_style(doc: Document, style_name: str):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def _set_run_font(
    run,
    font_name: str = WORD_EXPORT_FONT_LATIN,
    east_asia_font: str = WORD_EXPORT_FONT_EAST_ASIA,
    size: Pt | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _style_paragraph_runs(
    paragraph,
    font_name: str = WORD_EXPORT_FONT_LATIN,
    east_asia_font: str = WORD_EXPORT_FONT_EAST_ASIA,
) -> None:
    for run in paragraph.runs:
        _set_run_font(run, font_name=font_name, east_asia_font=east_asia_font)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_text_paragraph(doc: Document, text: str, style_name: str = "Normal"):
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.add_run(text)
    _style_paragraph_runs(paragraph)
    return paragraph


def _add_heading_paragraph(
    doc: Document,
    text: str,
    level: int,
    state: dict[str, bool],
    allow_title_style: bool = False,
) -> None:
    style_name = (
        "Title"
        if allow_title_style and level == 1 and not state["has_title"]
        else f"Heading {min(level, 6)}"
    )
    paragraph = _add_text_paragraph(doc, text, style_name=style_name)
    paragraph.paragraph_format.keep_with_next = True
    state["has_title"] = True


def _add_body_paragraph(doc: Document, text: str):
    return _add_text_paragraph(doc, text, style_name="Normal")


def _add_list_item(doc: Document, text: str, ordered: bool = False):
    style_name = "List Number" if ordered else "List Bullet"
    return _add_text_paragraph(doc, text, style_name=style_name)


def _add_caption(doc: Document, text: str):
    paragraph = _add_text_paragraph(doc, text, style_name="ChatGPT Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def _add_quote_block(doc: Document, text: str) -> None:
    paragraph = _add_text_paragraph(doc, text, style_name="ChatGPT Quote")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    _set_paragraph_shading(paragraph, WORD_EXPORT_QUOTE_BACKGROUND)


def _add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="ChatGPT Code")
    lines = text.rstrip("\n").splitlines() or [text]
    for index, line in enumerate(lines):
        run = paragraph.add_run(line if line else " ")
        _set_run_font(
            run,
            font_name=WORD_EXPORT_FONT_MONO,
            east_asia_font=WORD_EXPORT_FONT_MONO,
            size=Pt(9.5),
            color=WORD_EXPORT_TEXT_COLOR,
        )
        if index < len(lines) - 1:
            run.add_break()
    _set_paragraph_shading(paragraph, WORD_EXPORT_CODE_BACKGROUND)


def _add_table_from_html(doc: Document, table_tag: Tag) -> bool:
    row_tags = table_tag.find_all("tr")
    if not row_tags:
        return False

    rows: list[list[str]] = []
    header_rows: set[int] = set()
    column_count = 0

    for row_index, row_tag in enumerate(row_tags):
        cell_tags = row_tag.find_all(["th", "td"], recursive=False)
        if not cell_tags:
            cell_tags = row_tag.find_all(["th", "td"])
        if not cell_tags:
            continue
        row_values = [cell.get_text(" ", strip=True) for cell in cell_tags]
        rows.append(row_values)
        column_count = max(column_count, len(row_values))
        if any(cell.name == "th" for cell in cell_tags):
            header_rows.add(len(rows) - 1)

    if not rows or column_count == 0:
        return False

    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row_values in enumerate(rows):
        for column_index in range(column_count):
            text = row_values[column_index] if column_index < len(row_values) else ""
            cell = table.cell(row_index, column_index)
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                _style_paragraph_runs(paragraph)
                if row_index in header_rows:
                    for run in paragraph.runs:
                        run.font.bold = True
                if row_index in header_rows:
                    _set_cell_shading(cell, "EEF2F7")

    doc.add_paragraph("")
    return True


def _add_docx_image(doc: Document, image_buffer: io.BytesIO, caption_text: str | None = None) -> None:
    image_buffer.seek(0)
    doc.add_picture(image_buffer, width=Inches(WORD_EXPORT_IMAGE_WIDTH_INCHES))
    paragraph = doc.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4 if caption_text else 10)
    if caption_text:
        _add_caption(doc, caption_text)


def _decode_data_image_uri(image_src: str) -> io.BytesIO | None:
    if not image_src.startswith("data:image"):
        return None

    try:
        _, encoded = image_src.split(",", 1)
        image_bytes = base64.b64decode(encoded)
    except (ValueError, binascii.Error):
        return None

    return io.BytesIO(image_bytes)


def build_docx_from_html(html_text: str) -> bytes:
    html_text = normalize_trailing_punctuation_before_figure_placeholder(html_text)

    doc = Document()
    _configure_doc_style(doc)

    soup = BeautifulSoup(html_text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    root = soup.find("main") or soup.body or soup

    block_tags = {
        "article",
        "aside",
        "blockquote",
        "div",
        "figure",
        "figcaption",
        "main",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "li",
        "ul",
    }
    state = {"has_title": False}

    def add_image_from_tag(img_tag: Tag, caption_text: str | None = None) -> bool:
        image_src = (img_tag.get("src") or "").strip()
        image_buffer = _decode_data_image_uri(image_src)
        if image_buffer is None:
            return False
        _add_docx_image(doc, image_buffer, caption_text=caption_text)
        return True

    def walk(node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                parsed_lines = _split_markdown_heading_lines(text)
                if parsed_lines:
                    for line_kind, line_text in parsed_lines:
                        if line_kind == "heading":
                            _add_heading_paragraph(doc, line_text, level=1, state=state)
                        else:
                            _add_body_paragraph(doc, line_text)
                else:
                    _add_body_paragraph(doc, text)
            return

        if not isinstance(node, Tag):
            return

        if node.name in {"script", "style", "noscript"}:
            return

        if node.name == "br":
            return

        if node.name and node.name.startswith("h") and len(node.name) == 2 and node.name[1].isdigit():
            text = node.get_text(" ", strip=True)
            if text:
                _add_heading_paragraph(doc, text, level=1, state=state, allow_title_style=True)
            return

        if node.name == "li":
            text = node.get_text(" ", strip=True)
            if text:
                _add_list_item(doc, text, ordered=(node.parent is not None and node.parent.name == "ol"))
            return

        if node.name == "img":
            add_image_from_tag(node)
            return

        if node.name == "figure":
            img_tag = node.find("img")
            caption_tag = node.find("figcaption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else None
            if img_tag is not None:
                add_image_from_tag(img_tag, caption_text=caption_text)
            elif caption_text:
                _add_caption(doc, caption_text)
            return

        if node.name == "div" and "report-figure-block" in (node.get("class") or []):
            img_tag = node.find("img")
            caption_tag = node.find(class_="report-figure-caption")
            caption_text = caption_tag.get_text(" ", strip=True) if caption_tag else None
            if img_tag is not None:
                add_image_from_tag(img_tag, caption_text=caption_text)
            elif caption_text:
                _add_caption(doc, caption_text)
            return

        if node.name == "figcaption":
            text = node.get_text(" ", strip=True)
            if text:
                _add_caption(doc, text)
            return

        if node.name == "blockquote":
            text = node.get_text("\n", strip=True)
            if text:
                _add_quote_block(doc, text)
            return

        if node.name == "pre":
            text = node.get_text("\n", strip=False).strip("\n")
            if text:
                _add_code_block(doc, text)
            return

        if node.name == "table":
            _add_table_from_html(doc, node)
            return

        if node.name == "p":
            direct_images = [child for child in node.children if isinstance(child, Tag) and child.name == "img"]
            if direct_images and not node.get_text(" ", strip=True):
                for img_tag in direct_images:
                    add_image_from_tag(img_tag)
                return
            text = node.get_text("\n", strip=True)
            if text:
                parsed_lines = _split_markdown_heading_lines(text)
                if parsed_lines:
                    for line_kind, line_text in parsed_lines:
                        if line_kind == "heading":
                            _add_heading_paragraph(doc, line_text, level=1, state=state)
                        else:
                            _add_body_paragraph(doc, line_text)
                    return
                _add_body_paragraph(doc, text)
            return

        if node.name == "hr":
            doc.add_paragraph("")
            return

        if node.name in block_tags:
            direct_text = "\n".join(
                text.strip()
                for text in node.find_all(string=True, recursive=False)
                if text and text.strip()
            )
            handled_direct_text = False
            if direct_text and node.name not in {"ol", "ul"}:
                parsed_lines = _split_markdown_heading_lines(direct_text)
                if parsed_lines:
                    for line_kind, line_text in parsed_lines:
                        if line_kind == "heading":
                            _add_heading_paragraph(doc, line_text, level=1, state=state)
                        else:
                            _add_body_paragraph(doc, line_text)
                else:
                    _add_body_paragraph(doc, direct_text)
                handled_direct_text = True
            for child in node.children:
                if isinstance(child, Tag) and child.name == "img":
                    add_image_from_tag(child)
                elif isinstance(child, Tag):
                    walk(child)
                elif not handled_direct_text:
                    walk(child)
            return

        for child in node.children:
            walk(child)

    for child in root.children:
        walk(child)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def markdown_to_html(markdown_text: str, title: str = "Analysis Report") -> str:
    toc_items: list[tuple[str, str, int]] = []
    body_parts: list[str] = []
    in_ul = False

    def close_list() -> None:
        nonlocal in_ul
        if in_ul:
            body_parts.append("</ul>")
            in_ul = False

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            close_list()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            close_list()
            level = len(heading_match.group(1))
            text = html.escape(heading_match.group(2).strip())
            section_id = f"sec-{len(toc_items) + 1}"
            toc_items.append((section_id, text, level))
            body_parts.append(f"<h{level} id='{section_id}'>{text}</h{level}>")
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            if not in_ul:
                body_parts.append("<ul>")
                in_ul = True
            body_parts.append(f"<li>{html.escape(bullet_match.group(1).strip())}</li>")
            continue

        close_list()
        body_parts.append(f"<p>{html.escape(line)}</p>")

    close_list()

    toc_html = "".join(
        f"<a href='#{section_id}' class='toc-level-{level}'>{text}</a>"
        for section_id, text, level in toc_items
    )
    body_html = "\n".join(body_parts)

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    :root {{
      --bg: #f4f1ea;
      --panel: #fffaf2;
      --ink: #1f2937;
      --muted: #6b7280;
      --accent: #b45309;
      --border: #eadfce;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: Georgia, "Noto Serif SC", serif;
      color: var(--ink);
      background:
        radial-gradient(circle at top left, rgba(180, 83, 9, 0.08), transparent 28%),
        linear-gradient(180deg, #f8f4ed 0%, var(--bg) 100%);
    }}
    .layout {{
      display: grid;
      grid-template-columns: minmax(220px, 280px) minmax(0, 1fr);
      gap: 24px;
      max-width: 1280px;
      margin: 0 auto;
      padding: 24px;
    }}
    aside {{
      position: sticky;
      top: 24px;
      align-self: start;
      background: rgba(255, 250, 242, 0.92);
      border: 1px solid var(--border);
      border-radius: 20px;
      padding: 20px;
      backdrop-filter: blur(8px);
    }}
    aside h2 {{
      margin: 0 0 12px 0;
      font-size: 1.1rem;
    }}
    aside nav {{
      display: flex;
      flex-direction: column;
      gap: 8px;
    }}
    aside a {{
      color: var(--muted);
      text-decoration: none;
      line-height: 1.4;
    }}
    aside a:hover {{
      color: var(--accent);
    }}
    .toc-level-2 {{ padding-left: 12px; }}
    .toc-level-3 {{ padding-left: 24px; }}
    .toc-level-4, .toc-level-5, .toc-level-6 {{ padding-left: 36px; }}
    main {{
      background: rgba(255, 250, 242, 0.96);
      border: 1px solid var(--border);
      border-radius: 28px;
      padding: 40px;
      box-shadow: 0 16px 60px rgba(31, 41, 55, 0.08);
    }}
    h1, h2, h3, h4, h5, h6 {{
      color: #111827;
      line-height: 1.3;
      margin-top: 1.5em;
      margin-bottom: 0.6em;
    }}
    h1 {{ font-size: 2.2rem; }}
    h2 {{ font-size: 1.7rem; }}
    h3 {{ font-size: 1.35rem; }}
    p, li {{
      line-height: 1.8;
      font-size: 1rem;
    }}
    ul {{
      padding-left: 1.4rem;
      line-height: 1.8;
    }}
    @media (max-width: 900px) {{
      .layout {{
        grid-template-columns: 1fr;
      }}
      aside {{
        position: static;
      }}
      main {{
        padding: 24px;
      }}
    }}
  </style>
</head>
<body>
  <div class="layout">
    <aside>
      <h2>目录</h2>
      <nav>{toc_html}</nav>
    </aside>
    <main>{body_html}</main>
  </div>
</body>
</html>
"""


def build_docx_from_markdown(markdown_text: str) -> bytes:
    markdown_text = normalize_trailing_punctuation_before_figure_placeholder(markdown_text)

    doc = Document()
    _configure_doc_style(doc)
    state = {"has_title": False}
    code_lines: list[str] = []
    in_code_block = False

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped_line = line.strip()

        if stripped_line.startswith("```"):
            if in_code_block:
                _add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped_line:
            continue

        parsed_segments = _split_markdown_heading_lines(stripped_line)
        if parsed_segments:
            for line_kind, line_text in parsed_segments:
                if line_kind == "heading":
                    _add_heading_paragraph(doc, line_text, level=1, state=state)
                else:
                    _add_body_paragraph(doc, line_text)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped_line)
        if ordered_match:
            _add_list_item(doc, ordered_match.group(1).strip(), ordered=True)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped_line)
        if bullet_match:
            _add_list_item(doc, bullet_match.group(1).strip(), ordered=False)
            continue

        quote_match = re.match(r"^>\s?(.*)$", stripped_line)
        if quote_match:
            _add_quote_block(doc, quote_match.group(1).strip())
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\((.+)\)$", stripped_line)
        if image_match:
            image_buffer = _decode_data_image_uri(image_match.group(2).strip())
            if image_buffer is not None:
                caption_text = image_match.group(1).strip() or None
                _add_docx_image(doc, image_buffer, caption_text=caption_text)
            continue

        _add_body_paragraph(doc, stripped_line)

    if in_code_block and code_lines:
        _add_code_block(doc, "\n".join(code_lines))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
