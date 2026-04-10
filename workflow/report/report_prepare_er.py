import ast
import io
import re
from io import BytesIO

import streamlit as st
from tqdm import tqdm
from stqdm import stqdm
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.express as px
import plotly.io as pio
from concurrent.futures import ThreadPoolExecutor, as_completed

from utils.sanitize_code import sanitize_code
from workflow.report.report_core import *


def report_prepare(agents, parallel=True, max_workers=4):
    report_agent = agents[-1]
    toc = report_agent.load_outline()
    if toc is None:
        st.error("请先生成目录")
        return

    # 获取详细程度和用户要求（尝试调用 agent 的 load 方法，如果未定义则尝试直接访问属性）
    # 注意：这里假设 ReportAgent 有对应的 load_outline_length 和 load_user_input 方法
    # 如果没有，请根据您的 Agent 定义进行调整，或者使用 getattr(report_agent, 'outline_length', '')
    current_length = getattr(report_agent, 'load_outline_length', lambda: getattr(report_agent, 'outline_length', 'Standard'))()
    current_user_input = getattr(report_agent, 'load_user_input', lambda: getattr(report_agent, 'user_input', 'Default'))()
    
    # 构造当前的缓存 Key：(目录内容字符串, 详细程度, 用户要求)
    # 只要这三个没变，生成的文本内容理论上就不需要变
    current_cache_key = (str(toc), str(current_length), str(current_user_input))

    # === 检查缓存 ===
    # 从 report_agent 中获取上一次生成的参数 (monkey-patching 属性)
    last_cache_key = getattr(report_agent, '_last_gen_cache_key', None)
    existing_report = report_agent.load_report()

    # 如果 Key 一致，且内存中已经有 Reportcore 对象（且不为空），则跳过生成
    if last_cache_key == current_cache_key and existing_report is not None:
        # 简单检查一下 root 下是否有内容
        if hasattr(existing_report, 'root') and existing_report.root.children:
            st.success(f"⚡ 检测到报告要求与大纲未变更，直接复用现有文档结构进行格式转换。")
            return

    # ==========================================
    # 以下是原有的生成逻辑 (LLM 处理部分)
    # ==========================================

    toc = sanitize_code(toc)

    # === 汇总各分析模块的摘要 ===
    agent_abstracts = {}
    with st.spinner("正在汇总各分析模块的结果..."):
        for i in stqdm(range(len(agents) - 1)):
            agent_abstracts[i] = agents[i].check_abstract()

    # === 更新 toc 的 FIG 列表 ===
    selected_full_contents_vis = agents[2].check_full()
    toc = report_agent.selected_photo_update_toc(toc, selected_full_contents_vis)
    toc = sanitize_code(toc)
    # print(toc)
    try:
        toc = ast.literal_eval(toc)
    except Exception:
        pass

    # === 更新 toc 的 模块选择 列表 ===
    with st.spinner("正在匹配各章节所需的分析模块..."):
        toc_with_choice = report_agent.update_toc_with_relevant_sections(toc, agent_abstracts)
        toc_with_choice = sanitize_code(toc_with_choice)
        try:
            toc_with_choice = ast.literal_eval(toc_with_choice)
        except Exception:
            pass

    # === 初始化报告结构 ===
    doc = Reportcore()
    doc.add_heading('数据分析报告', 0)
    
    # 确保 session_state 能够被线程访问
    if 'selected_model' in st.session_state:
        selected_model = st.session_state.selected_model
    else:
        selected_model = "default" # Fallback

    def process_section(idx, t, t_w_c, history_content=""):
        # 线程内可能丢失 session_state，重新赋值
        st.session_state.selected_model = selected_model
        # t: ('标题', 层级, 内容大纲, [figs], [modules])
        _, _, _, _, choice_list = t_w_c
        selected_full_contents = {i: agents[i].check_full() for i in choice_list if i < len(agents) - 1}
        content = report_agent.write_section_body(toc, t, selected_full_contents, history_content)
        # print(idx)
        return (idx, t, content)

    results = []

    # 串行或并行
    if not parallel:
        with st.spinner("正在串行生成各章节内容（带上下文）..."):
            history_content = ""
            for idx, t in stqdm(enumerate(toc)):
                t_w_c = toc_with_choice[idx]
                _, _, content = process_section(idx, t, t_w_c, history_content)
                results.append((idx, t, content))
                history_content += f"\n\n{t[0]}\n{content}"
    else:
        with st.spinner(f"正在并行生成各章节内容（{max_workers}线程）..."):
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # print(toc_with_choice)
                futures = {
                    executor.submit(process_section, idx, t, toc_with_choice[idx], ""): idx
                    for idx, t in enumerate(toc)
                }
                for future in stqdm(as_completed(futures), total=len(futures)):
                    try:
                        results.append(future.result())
                    except Exception as e:
                        print(f"章节生成失败: {e}")

    # 排序 & 写入报告
    results.sort(key=lambda x: x[0])
    for _, t, content in results:
        doc.add_heading(t[0], level=t[1])
        doc.add_paragraph(content)

    report_agent.save_report(doc)
    
    # === 关键步骤：更新缓存状态 ===
    # 只有在完整生成成功后，才更新 last_gen_cache_key
    report_agent._last_gen_cache_key = current_cache_key